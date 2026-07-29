import io
import re
import docx
import docx2txt

def parse_docx(file_content: bytes, filename: str = "") -> str:
    """
    Parses DOCX/DOC documents.
    1. Uses python-docx for paragraphs and tables.
    2. Falls back to docx2txt.
    3. For binary legacy DOC files, uses regex text extraction fallback.
    """
    is_legacy_doc = filename.lower().endswith('.doc') or file_content.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1')

    if not is_legacy_doc:
        try:
            # Try python-docx
            doc = docx.Document(io.BytesIO(file_content))
            full_text = []
            
            # Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and (not full_text or full_text[-1] != cell_text):
                            full_text.append(cell_text)
                            
            return "\n".join(full_text)
        except Exception:
            # Fallback to docx2txt
            try:
                text = docx2txt.process(io.BytesIO(file_content))
                if text and text.strip():
                    return text
            except Exception:
                pass

    # Regex string extraction fallback for legacy binary .doc or corrupted files
    try:
        # Extract sequences of 4 or more printable ASCII/UTF-8 chars
        pattern = re.compile(b'[a-zA-Z0-9\\s\\.,;:!\\?\\-\\@\\_\\(\\)\\[\\]\\{\\}/\\\\"\'\\&\\+\\#\\*\\=\\%\\<\\>]{4,}')
        text_runs = []
        for match in pattern.finditer(file_content):
            try:
                decoded = match.group(0).decode('utf-8', errors='ignore').strip()
                # Filter out pure noise / spacing
                if decoded and len(re.sub(r'[\s\.\,\-_]+', '', decoded)) > 2:
                    text_runs.append(decoded)
            except Exception:
                pass
        
        extracted_text = "\n".join(text_runs)
        # Normalize double whitespace/newlines
        extracted_text = re.sub(r'\n+', '\n', extracted_text)
        extracted_text = re.sub(r'[ \t]+', ' ', extracted_text)
        
        if len(extracted_text.strip()) > 50:
            return extracted_text.strip()
    except Exception as e:
        raise ValueError(f"Regex extraction fallback failed: {str(e)}")

    raise ValueError(
        "Could not parse Word document. If this is a legacy .doc file, "
        "please save it as a modern .docx file and try uploading again."
    )
